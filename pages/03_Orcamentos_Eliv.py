# pages/03_Orcamentos_Eliv.py
import io
import datetime as dt
import streamlit as st
from docx import Document

# --------------- CONFIG ----------------
st.set_page_config(page_title="Orçamentos Eliv", page_icon="📦")
st.title("Orçamentos Eliv")

K = "eliv_"  # prefixo para keys únicas

# Tabela de preços de lista dos pacotes
PACOTES = {
    "Básico":   {"mensal": 349.00, "avista": 1884.60},
    "Especial": {"mensal": 359.00, "avista": 1938.60},
    "Premium":  {"mensal": 499.00, "avista": 2694.60},
}

def br_money(v: float) -> str:
    s = f"R${v:,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")

def docx_orcamento_universidade(
    pacote_nome, forma_pag, preco_lista_total, preco_desconto_total, mensal_final,
    desconto_pac_pct,
    universidade, contato, autor, obra, paginas,
    capa="flexível", papel="offset 75g", miolo="preto e branco",
    preco_capa=99.90, tiragem_qtd=100, desc_tiragem_pct=30.0,
    ebook_preco=60.90
) -> bytes:
    doc = Document()
    h = doc.add_heading("Orçamento – Publicação ELIV", level=1)

    data = dt.date.today().strftime("%d/%m/%Y")
    p = doc.add_paragraph()
    p.add_run(f"Data do orçamento: {data}\n").bold = True

    doc.add_paragraph(
        f'A pedido de à {contato}, {universidade}, {autor}\n'
        f'Editoração de obra "{obra}" com extensão de {paginas} páginas, '
        f'com selo acadêmico, capa {capa}, papel {papel}, impressão do miolo {miolo}.'
    )

    doc.add_paragraph("")
    doc.add_paragraph(f"Pacote selecionado: {pacote_nome}")
    doc.add_paragraph(f"Forma de pagamento: {forma_pag}")
    doc.add_paragraph(
        f"Valor de lista: {br_money(preco_lista_total)} | "
        f"Desconto aplicado: {desconto_pac_pct:.0f}% | "
        f"Valor com desconto: {br_money(preco_desconto_total)}"
        + (f" (6x de {br_money(mensal_final)})" if mensal_final else "")
    ).bold = True

    doc.add_paragraph("")
    doc.add_paragraph("Serviços inclusos:")
    items = [
        "Diagramação",
        "Design de capa e miolo",
        "Registro ISBN e DOI",
        "Produção das versões de Livro Físico e E-book",
        "Venda",
        "Distribuição virtual",
    ]
    for it in items:
        doc.add_paragraph(f"• {it}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Preço de capa sugerido (livro físico): {br_money(preco_capa)}")
    doc.add_paragraph(
        f"Para {tiragem_qtd} exemplares, desconto de {desc_tiragem_pct:.0f}%: "
        f"preço unitário {br_money(preco_capa * (1 - desc_tiragem_pct/100))}"
    )
    total_tiragem = preco_capa * (1 - desc_tiragem_pct/100) * tiragem_qtd
    doc.add_paragraph(f"Total da tiragem ({tiragem_qtd} un.): {br_money(total_tiragem)}").bold = True

    doc.add_paragraph(f"Preço do e-book: {br_money(ebook_preco)}.")
    total_geral = preco_desconto_total + total_tiragem
    doc.add_paragraph(
        f"Custo total: Editoração ({br_money(preco_desconto_total)}) + "
        f"Tiragem ({br_money(total_tiragem)}): {br_money(total_geral)}"
    ).bold = True

    doc.add_paragraph("")
    doc.add_paragraph(
        "*O preço de capa é uma estimativa e pode mudar após a diagramação caso haja alteração das especificações."
    )
    doc.add_paragraph(
        "*A tiragem deverá ser contratada ao final do processo de publicação, caso o autor deseje."
    )
    doc.add_paragraph("Validade da proposta: 15 dias.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def docx_orcamento_comum(
    pacote_nome, forma_pag, preco_lista_total, preco_desconto_total, mensal_final,
    desconto_pac_pct, cliente, consultor, obra, paginas,
    observacoes=""
) -> bytes:
    doc = Document()
    doc.add_heading("Orçamento – Publicação ELIV", level=1)

    data = dt.date.today().strftime("%d/%m/%Y")
    doc.add_paragraph(f"Data do orçamento: {data}")
    doc.add_paragraph(f"Cliente: {cliente or '—'}")
    doc.add_paragraph(f"Consultor: {consultor or '—'}")

    doc.add_paragraph("")
    doc.add_paragraph(f'Obra: "{obra}" – {paginas} páginas.')
    doc.add_paragraph(f"Pacote selecionado: {pacote_nome}")
    doc.add_paragraph(f"Forma de pagamento: {forma_pag}")
    doc.add_paragraph(
        f"Valor de lista: {br_money(preco_lista_total)} | "
        f"Desconto aplicado: {desconto_pac_pct:.0f}% | "
        f"Valor com desconto: {br_money(preco_desconto_total)}"
        + (f" (6x de {br_money(mensal_final)})" if mensal_final else "")
    ).bold = True

    if observacoes:
        doc.add_paragraph("")
        doc.add_paragraph("Observações:")
        doc.add_paragraph(observacoes)

    doc.add_paragraph("")
    doc.add_paragraph("Validade da proposta: 15 dias.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ----------------- UI -------------------
c0, c1 = st.columns(2)
cliente   = c0.text_input("Nome do cliente", key=f"{K}cliente")
consultor = c1.text_input("Consultor", key=f"{K}consultor")

c2, c3 = st.columns(2)
obra     = c2.text_input("Título da obra", key=f"{K}obra")
paginas  = c3.number_input("Número de páginas", min_value=10, step=10, value=250, key=f"{K}pag")

st.divider()

c4, c5, c6 = st.columns(3)
pacote = c4.selectbox("Pacote", list(PACOTES.keys()), index=1, key=f"{K}pacote")
desconto_pac_pct = c5.number_input("% de desconto no pacote", min_value=0.0, max_value=100.0, value=15.0, step=1.0, key=f"{K}descpac")
forma_pag = c6.radio("Forma de pagamento", ["6x sem juros", "À vista (PIX)"], horizontal=True, key=f"{K}fpag")

# Preços de lista do pacote
precos = PACOTES[pacote]
preco_lista_total_6x = precos["mensal"] * 6
preco_lista_avista   = precos["avista"]

if forma_pag.startswith("6x"):
    preco_lista_escolhido = preco_lista_total_6x
else:
    preco_lista_escolhido = preco_lista_avista

preco_desconto_total = preco_lista_escolhido * (1 - desconto_pac_pct/100.0)
mensal_final = (preco_desconto_total / 6.0) if forma_pag.startswith("6x") else None

st.caption(
    f"Lista — 6x: {br_money(preco_lista_total_6x)} (6×{br_money(precos['mensal'])}) · "
    f"À vista: {br_money(preco_lista_avista)}"
)
st.success(
    ("Pacote " + pacote + " | " + forma_pag +
     f" | Com desconto: {br_money(preco_desconto_total)}" +
     (f" (6x de {br_money(mensal_final)})" if mensal_final else "")
    )
)

st.divider()

# --------- Modo Universidade (mostra tiragem) ----------
modo_uni = st.toggle("Orçamento para Universidade", value=False, key=f"{K}uni_toggle")

if modo_uni:
    u1, u2 = st.columns(2)
    universidade = u1.text_input("Universidade", key=f"{K}uni_nome")
    contato      = u2.text_input("Contato (Profa./Prof., telefone, etc.)", key=f"{K}uni_contato")

    st.subheader("Tiragem da universidade")
    t1, t2, t3 = st.columns(3)
    preco_capa = t1.number_input("Preço de capa (R$)", min_value=0.0, value=99.90, step=0.10, format="%.2f", key=f"{K}capafis")
    tiragem_qtd = t2.number_input("Quantidade", min_value=1, value=100, step=50, key=f"{K}tir_qtd")
    desc_tiragem_pct = t3.number_input("% de desconto na tiragem", min_value=0.0, max_value=100.0, value=30.0, step=1.0, key=f"{K}tir_desc")

    ebook_preco = st.number_input("Preço do e-book (R$)", min_value=0.0, value=60.90, step=0.10, format="%.2f", key=f"{K}ebookp")

    preco_unit_liquido = preco_capa * (1 - desc_tiragem_pct/100)
    total_tiragem = preco_unit_liquido * tiragem_qtd

    st.info(
        f"Unitário com desconto: {br_money(preco_unit_liquido)} | "
        f"Total tiragem ({tiragem_qtd}): {br_money(total_tiragem)}"
    )

st.divider()

# ---------------- Script pronto (sem emojis) ----------------
st.subheader("Script para WhatsApp/CRM")
data = dt.date.today().strftime("%d/%m/%Y")
if modo_uni:
    script = f"""Orçamento ELIV — {data}

Cliente/Instituição: {cliente or '—'}
Contato: {consultor or '—'}
Universidade: {st.session_state.get(f"{K}uni_nome",'—')}
Responsável na universidade: {st.session_state.get(f"{K}uni_contato",'—')}

Obra: "{obra or '—'}" — {int(paginas)} páginas.
Pacote: {pacote} | {forma_pag}
Preço de lista: {br_money(preco_lista_escolhido)}
Desconto no pacote: {desconto_pac_pct:.0f}% → {br_money(preco_desconto_total)}{(' (6x de ' + br_money(mensal_final) + ')') if mensal_final else ''}

Tiragem: {int(st.session_state.get(f'{K}tir_qtd',0))} un.
Preço de capa: {br_money(st.session_state.get(f'{K}capafis',0.0))}
Desconto para universidade: {st.session_state.get(f'{K}tir_desc',0.0):.0f}% → unitário {br_money(preco_unit_liquido)}
Total da tiragem: {br_money(total_tiragem)}

Preço do e-book: {br_money(st.session_state.get(f'{K}ebookp',0.0))}

Validade: 15 dias."""
else:
    script = f"""Orçamento ELIV — {data}

Cliente: {cliente or '—'}
Consultor: {consultor or '—'}
Obra: "{obra or '—'}" — {int(paginas)} páginas.

Pacote: {pacote} | {forma_pag}
Preço de lista: {br_money(preco_lista_escolhido)}
Desconto no pacote: {desconto_pac_pct:.0f}% → {br_money(preco_desconto_total)}{(' (6x de ' + br_money(mensal_final) + ')') if mensal_final else ''}

Validade: 15 dias."""
st.text_area("Script pronto para copiar", script, height=220, key=f"{K}script_area")
st.download_button("Baixar script (.txt)", data=script, file_name="orcamento_eliv.txt", mime="text/plain", key=f"{K}down_script")

st.divider()

# ---------------- DOCX ----------------
st.subheader("Gerar DOCX")

if modo_uni:
    if st.button("Gerar DOCX – Universidade", key=f"{K}btn_docx_uni"):
        bin_doc = docx_orcamento_universidade(
            pacote_nome=pacote,
            forma_pag=forma_pag,
            preco_lista_total=preco_lista_escolhido,
            preco_desconto_total=preco_desconto_total,
            mensal_final=mensal_final,
            desconto_pac_pct=desconto_pac_pct,
            universidade=st.session_state.get(f"{K}uni_nome",""),
            contato=st.session_state.get(f"{K}uni_contato",""),
            autor=consultor or "",
            obra=obra or "",
            paginas=int(paginas),
            preco_capa=st.session_state.get(f"{K}capafis",99.90),
            tiragem_qtd=int(st.session_state.get(f"{K}tir_qtd",0)),
            desc_tiragem_pct=st.session_state.get(f"{K}tir_desc",0.0),
            ebook_preco=st.session_state.get(f"{K}ebookp",0.0)
        )
        st.download_button("Baixar DOCX – Universidade", data=bin_doc, file_name="Orcamento_Universidade_ELIV.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"{K}down_uni")
else:
    obs = st.text_area("Observações (opcional)", key=f"{K}obs")
    if st.button("Gerar DOCX – Comum", key=f"{K}btn_docx_comum"):
        bin_doc = docx_orcamento_comum(
            pacote_nome=pacote,
            forma_pag=forma_pag,
            preco_lista_total=preco_lista_escolhido,
            preco_desconto_total=preco_desconto_total,
            mensal_final=mensal_final,
            desconto_pac_pct=desconto_pac_pct,
            cliente=cliente or "",
            consultor=consultor or "",
            obra=obra or "",
            paginas=int(paginas),
            observacoes=obs or ""
        )
        st.download_button("Baixar DOCX – Comum", data=bin_doc, file_name="Orcamento_Comum_ELIV.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"{K}down_comum")
